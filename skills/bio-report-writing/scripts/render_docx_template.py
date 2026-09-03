#!/usr/bin/env python3
"""Fill a DOCX-first report template with explicit, evidence-bound slots.

This is a small renderer helper for module R/Python report coders.  It never
calculates statistics or invents prose.  Without ``--final`` it deliberately
returns ``DRAFT``; a release caller must still run the contract and visual
validators.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
import hashlib
import json
import os
from pathlib import Path
import re
import sys
import tempfile

from validate_docx_structure import TEMPLATE_MARKER, _marker_bookmark_name, extract_template_markers, inspect
from validate_report_contract import (
    normalise_relative,
    resolve_under,
    validate_pack,
    validate_plan,
    validate_slot_alignment,
)


TABLE_SPECS = {
    "RESULTS": {
        "rows_key": "RESULT_TABLE_ROWS",
        "keys": ("name", "value", "unit", "source"),
        "caption_key": "TABLE:RESULTS.CAPTION",
    },
    "OUTPUTS": {
        "rows_key": "OUTPUT_TABLE_ROWS",
        "keys": ("path", "kind", "description", "purpose", "consumers"),
        "caption_key": "TABLE:OUTPUTS.CAPTION",
    },
    "VERSIONS": {
        "rows_key": "VERSION_TABLE_ROWS",
        "keys": ("name", "version", "purpose"),
        "caption_key": "TABLE:VERSIONS.CAPTION",
    },
}
SOURCE_SUFFIX = ".SOURCE"
# 这些词不是事实；release 中若进入 Note，说明 renderer 收到了泛化占位。
GENERIC_NOTE_TEXT = re.compile(
    r"(?:按(?:本次|当前)?配置|以原图为准|以图中为准|见原图|根据需要|待确认|待补(?:充)?|占位)",
    re.I,
)
QUESTION_TEXT = re.compile(r"(?:[？?]|哪些|如何|是否|为什么|什么|哪种|请判断)")
VALUE_PLACEHOLDER = re.compile(
    r"(?:\[\[[^\]]+\]\]|\{\{[^}]+\}\}|\b(?:TODO|TBD|REPLACE|EVIDENCE_REQUIRED|EVIDENCE_NEEDED)\b)",
    re.I,
)
VISIBLE_SOURCE_URL = re.compile(r"(?:https?://|\bdoi\s*:\s*10\.)", re.I)


def _load(path: Path) -> object:
    return json.loads(path.read_text(encoding="utf-8"))


def _scalar(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, (str, int, float)):
        return str(value)
    if isinstance(value, list):
        return "；".join(_scalar(item) for item in value)
    if isinstance(value, dict):
        return "；".join(f"{key}：{_scalar(item)}" for key, item in value.items())
    return str(value)


def _slot_text(value: object) -> str:
    if isinstance(value, dict):
        for key in ("text", "content", "body", "value", "description"):
            if value.get(key) not in (None, "", [], {}):
                return _scalar(value[key])
    return _scalar(value)


def _visible_references(value: object) -> str:
    """只显示来源名称、版本和用途；地址留在内部 provenance。"""
    items = value if isinstance(value, list) else [value]
    rendered: list[str] = []
    for item in items:
        if isinstance(item, dict):
            name = item.get("name", item.get("title", "来源"))
            version = item.get("version")
            purpose = item.get("purpose")
            text = _scalar(name)
            if version not in (None, ""):
                text += f"（版本 {_scalar(version)}"
                if purpose not in (None, ""):
                    text += f"；用途 {_scalar(purpose)}"
                text += "）"
            elif purpose not in (None, ""):
                text += f"（用途 {_scalar(purpose)}）"
            rendered.append(text)
        elif item not in (None, ""):
            rendered.append(_scalar(item))
    return "；".join(rendered)


def _key(value: object) -> str:
    return str(value).strip().strip("[]").upper()


def _flatten_values(value: object) -> dict[str, object]:
    if not isinstance(value, dict):
        raise ValueError("slot values must be a JSON object")
    source = value.get("slots") if isinstance(value.get("slots"), dict) else value
    values = {_key(name): item for name, item in source.items() if name not in {"tables", "figures", "notes", "text", "_NOTE_ORDER"}}
    text = source.get("text") if isinstance(source, dict) else None
    if isinstance(text, dict):
        values.update({_key(name): item for name, item in text.items()})
    tables = source.get("tables") if isinstance(source, dict) else None
    if isinstance(tables, dict):
        for name, table in tables.items():
            kind = str(name).upper().replace("TABLE:", "")
            values[f"TABLE:{kind}"] = table
            if kind in TABLE_SPECS:
                values[TABLE_SPECS[kind]["rows_key"]] = table
                if isinstance(table, dict) and table.get("caption") not in (None, ""):
                    values[TABLE_SPECS[kind]["caption_key"]] = table["caption"]
    figures = source.get("figures") if isinstance(source, dict) else None
    if isinstance(figures, list):
        for index, figure in enumerate(figures, 1):
            if not isinstance(figure, dict):
                continue
            ident = str(figure.get("id", f"F{index}")).upper()
            prefix = f"FIGURE:{ident}"
            for suffix, key in (("TITLE", "title"), ("SOURCE", "path"), ("CAPTION", "caption"), ("CAPTION_FIELDS", "caption_fields")):
                if key in figure:
                    values[f"{prefix}.{suffix}"] = figure[key]
    notes = source.get("notes") if isinstance(source, dict) else None
    if isinstance(notes, list):
        ordered_notes: list[object] = []
        for index, note in enumerate(notes):
            if isinstance(note, dict):
                ident = str(note.get("id", "direction")).upper()
                values[f"NOTE:{ident}"] = note
                ordered_notes.append(note)
        if ordered_notes:
            values["_NOTE_ORDER"] = ordered_notes
    defaults = {
        "TABLE:RESULTS.CAPTION": "主要结果",
        "TABLE:OUTPUTS.CAPTION": "公开业务文件",
        "TABLE:VERSIONS.CAPTION": "软件与资源版本",
    }
    for name, default in defaults.items():
        values.setdefault(name, default)
    return values


def _lookup(values: dict[str, object], marker: str) -> object:
    name = _key(marker[2:-2] if marker.startswith("[[") else marker)
    if name in values:
        return values[name]
    # Permit dotted/colon aliases supplied by a renderer.
    aliases = {
        name.replace(".", ":"),
        name.replace(":", "."),
        name.replace(".", "_"),
    }
    # The pack/slot examples use both short names and the explicit report
    # prefix.  Keep this aliasing deterministic rather than accepting arbitrary
    # fuzzy matches.
    if name.startswith("REPORT_"):
        aliases.add(name[7:])
    elif name in {"SUMMARY", "SCOPE", "METHOD", "QC", "RESULT", "CONCLUSION", "LIMITATIONS"}:
        aliases.add(f"REPORT_{name}" if name == "SUMMARY" else f"ANALYSIS_{name}")
    for alias in aliases:
        if alias in values:
            return values[alias]
    return None


def values_from_pack(pack: dict[str, object]) -> dict[str, object]:
    """Build conservative defaults from the shared evidence pack.

    Module renderers should normally pass their own explicit slot map; these
    defaults only make the DOCX template useful for a one-off run.
    """
    points = [item for item in pack.get("analysis_points", []) if isinstance(item, dict)]
    point = points[0] if points else {}
    comparison = point.get("comparison") if isinstance(point.get("comparison"), dict) else {}
    results = point.get("results", [])
    outputs = [
            item for item in point.get("outputs", [])
            if isinstance(item, dict) and item.get("published")
        ]
    values: dict[str, object] = {
        "REPORT_TITLE": pack.get("title", "分析报告"),
        "REPORT_AUDIENCE": pack.get("audience", ""),
        "REPORT_SUMMARY": point.get("interpretation", ""),
        "ANALYSIS_SCOPE": point.get("scope", ""),
        "ANALYSIS_METHOD": point.get("method", ""),
        "ANALYSIS_QC": point.get("qc", ""),
        "QC.CONDITION": bool(point.get("qc")),
        "ANALYSIS_RESULT": results,
        "RESULTS.CONDITION": bool(results),
        "ANALYSIS_CONCLUSION": point.get("interpretation", ""),
        "ANALYSIS_LIMITATIONS": point.get("limitations", []),
        "OUTPUTS_INTRO": "当前分析实际发布的业务文件及其用途。",
        "REFERENCES": pack.get("references", []),
        "RESULT_TABLE_ROWS": results,
        "OUTPUT_TABLE_ROWS": outputs,
        "VERSION_TABLE_ROWS": pack.get("versions", []),
        "FIGURES.CONDITION": bool(point.get("figure_table_refs")),
    }
    # Note 只接受 pack 明确声明的真实 notes[]；不要由 comparison 自动
    # 合成泛化提示，否则无方向/单位证据时会伪造读者可见口径。
    notes = point.get("notes", pack.get("notes", []))
    if isinstance(notes, list):
        ordered_notes: list[object] = []
        for index, note in enumerate(notes):
            if isinstance(note, dict):
                note_id = str(note.get("id", "direction")).upper()
                values[f"NOTE:{note_id}"] = note
                ordered_notes.append(note)
        if ordered_notes:
            values["_NOTE_ORDER"] = ordered_notes
    for ref in point.get("figure_table_refs", []):
        if not isinstance(ref, dict):
            continue
        ident = str(ref.get("id", "F1"))
        prefix = f"FIGURE.{ident.upper()}"
        values[f"{prefix}.SOURCE"] = ref.get("path", "")
        values[f"{prefix}.TITLE"] = ref.get("title", ref.get("id", ident))
        values[f"{prefix}.CAPTION"] = ref.get("caption", "")
        values[f"{prefix}.CAPTION_FIELDS"] = ref.get("caption_fields", "")
    for table_name, spec in TABLE_SPECS.items():
        captions = {"RESULTS": "主要结果", "OUTPUTS": "公开业务文件", "VERSIONS": "软件与资源版本"}
        values.setdefault(spec["caption_key"], captions[table_name])
        values.setdefault(f"TABLE:{table_name}", "")
    return values


def _iter_paragraphs(document):
    # ``python-docx`` can create a short-lived proxy for the same XML paragraph
    # on each property access; using ``id(proxy)`` therefore causes IDs to be
    # reused and silently skips table-cell paragraphs.  Keep the underlying
    # lxml element itself in the set instead.
    seen: set[object] = set()

    def cells(table):
        for row in table.rows:
            for cell in row.cells:
                yield cell
                for nested in cell.tables:
                    yield from cells(nested)

    def add(container):
        for paragraph in getattr(container, "paragraphs", []):
            if paragraph._p not in seen:
                seen.add(paragraph._p)
                yield paragraph
        for table in getattr(container, "tables", []):
            for cell in cells(table):
                for paragraph in cell.paragraphs:
                    if paragraph._p not in seen:
                        seen.add(paragraph._p)
                        yield paragraph

    yield from add(document)
    for section in document.sections:
        for container in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            yield from add(container)


def _paragraph_text(paragraph) -> str:
    return "".join(run.text or "" for run in paragraph.runs)


def _replace_span(paragraph, start: int, end: int, replacement: str) -> None:
    """Replace a text span even when Word split a marker across runs."""
    runs = list(paragraph.runs)
    positions: list[tuple[int, int]] = []
    cursor = 0
    for run in runs:
        text = run.text or ""
        positions.append((cursor, cursor + len(text)))
        cursor += len(text)
    touched = [index for index, (left, right) in enumerate(positions) if right > start and left < end]
    if not touched:
        return
    first = touched[0]
    last = touched[-1]
    left_offset = start - positions[first][0]
    right_offset = end - positions[last][0]
    first_text = runs[first].text or ""
    if first == last:
        runs[first].text = first_text[:left_offset] + replacement + first_text[right_offset:]
        return
    last_text = runs[last].text or ""
    suffix = last_text[right_offset:]
    runs[first].text = first_text[:left_offset] + replacement
    for index in touched[1:]:
        runs[index].text = suffix if index == last else ""


def _figure_body_size_emu(paragraph) -> tuple[int, int]:
    """按模板页面版心计算图片最大宽高（EMU），保留源图宽高比。"""
    section = paragraph.part.document.sections[0]
    width = int(section.page_width - section.left_margin - section.right_margin)
    height = int(section.page_height - section.top_margin - section.bottom_margin)
    left = paragraph.paragraph_format.left_indent
    right = paragraph.paragraph_format.right_indent
    if left is not None:
        width -= int(left)
    if right is not None:
        width -= int(right)
    return max(width, 1), max(height, 1)


def _clear_figure_paragraph_box(paragraph) -> None:
    """清除旧模板图段落的蓝框、底色和缩进；蓝框只属于 Note。"""
    from docx.oxml.ns import qn
    ppr = paragraph._p.get_or_add_pPr()
    for name in ("pBdr", "shd", "ind"):
        node = ppr.find(qn(f"w:{name}"))
        if node is not None:
            ppr.remove(node)


def _clear_all_figure_paragraph_boxes(document) -> None:
    """兼容旧模板：已有 drawing 的段落一律去掉 Note 装饰属性。"""
    drawing_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    for paragraph in _iter_paragraphs(document):
        if paragraph._p.find(f".//{drawing_tag}") is not None:
            _clear_figure_paragraph_box(paragraph)


def _remove_drawings(paragraph) -> None:
    """替换旧图时先移除同一 slot 的已有 drawing，避免一图多嵌。"""
    drawing_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    for drawing in list(paragraph._p.findall(f".//{drawing_tag}")):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)


def _insert_picture(paragraph, marker: str, value: object, root: Path | None) -> bool:
    """插入真实源图并按模板版心等比缩放，不裁切、不分幅。"""
    _clear_figure_paragraph_box(paragraph)
    _remove_drawings(paragraph)
    raw = _scalar(value)
    if not raw:
        return False
    candidate = Path(raw)
    if root is not None and not candidate.is_absolute():
        relative = normalise_relative(raw)
        candidate = resolve_under(root, relative) if relative else None
    else:
        candidate = candidate.resolve()
    if candidate is None or not candidate.is_file() or candidate.suffix.lower() not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
        return False
    try:
        from PIL import Image
        with Image.open(candidate) as image:
            source_width, source_height = image.size
    except Exception:
        return False
    if source_width <= 0 or source_height <= 0:
        return False
    runs = list(paragraph.runs)
    full = _paragraph_text(paragraph)
    start = full.find(marker)
    if start < 0:
        return False
    end = start + len(marker)
    positions: list[tuple[int, int]] = []
    cursor = 0
    for run in runs:
        text = run.text or ""
        positions.append((cursor, cursor + len(text)))
        cursor += len(text)
    touched = [index for index, (left, right) in enumerate(positions) if right > start and left < end]
    if not touched:
        return False
    first, last = touched[0], touched[-1]
    left_offset = start - positions[first][0]
    right_offset = end - positions[last][0]
    first_text = runs[first].text or ""
    suffix = (runs[last].text or "")[right_offset:]
    runs[first].text = first_text[:left_offset]
    for index in touched[1:]:
        runs[index].text = ""
    # Explicit width/height prevents python-docx from using a large native DPI
    # size that can exceed the page body.  Height is derived from the source
    # ratio; no crop rectangle or panel slicing is introduced.
    from docx.shared import Emu
    max_width_emu, max_height_emu = _figure_body_size_emu(paragraph)
    scale = min(
        1.0,
        max_width_emu / source_width,
        max_height_emu / max(1, round(max_width_emu * source_height / source_width)),
    )
    width_emu = max(1, round(max_width_emu * scale))
    height_emu = max(1, round(width_emu * source_height / source_width))
    runs[first].add_picture(str(candidate), width=Emu(width_emu), height=Emu(height_emu))
    if suffix:
        paragraph.add_run(suffix)
    return True


def _cell_set_text(cell, value: object) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(_scalar(value))
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _table_kind(table) -> str | None:
    text = "".join(cell.text for row in table.rows for cell in row.cells).upper()
    marker_prefixes = {
        "RESULTS": ("[[RESULT.NAME]]", "[[RESULT.VALUE]]", "[[RESULT.UNIT]]", "[[RESULT.SOURCE]]"),
        "OUTPUTS": ("[[OUTPUT.PATH]]", "[[OUTPUT.KIND]]", "[[OUTPUT.DESCRIPTION]]", "[[OUTPUT.PURPOSE]]", "[[OUTPUT.CONSUMERS]]"),
        "VERSIONS": ("[[VERSION.NAME]]", "[[VERSION.VALUE]]", "[[VERSION.PURPOSE]]"),
    }
    for kind, markers in marker_prefixes.items():
        if any(marker in text for marker in markers):
            return kind
    return None


def _records_and_columns(raw: object, kind: str) -> tuple[list[object], list[str] | None]:
    if isinstance(raw, dict):
        records = raw.get("rows", raw.get("data", []))
        columns = raw.get("columns")
        if isinstance(columns, list):
            names: list[str] = []
            for item in columns:
                if isinstance(item, dict):
                    names.append(str(item.get("key", item.get("name", ""))))
                else:
                    names.append(str(item))
            columns = names
        else:
            columns = None
    else:
        records, columns = raw, None
    if not isinstance(records, list):
        records = []
    return records, columns


def _filter_public_output_rows(values: dict[str, object], report_file: str | None = None) -> None:
    """输出表只保留有消费者的业务文件，不展示报告和控制面文件。"""
    raw = values.get("OUTPUT_TABLE_ROWS")
    if raw is None:
        raw = values.get("TABLE:OUTPUTS")
    if isinstance(raw, dict) and isinstance(raw.get("rows", raw.get("data")), list):
        rows = raw.get("rows", raw.get("data"))
        container = dict(raw)
        key = "rows" if "rows" in raw else "data"
    elif isinstance(raw, list):
        rows = raw
        container = None
        key = "rows"
    else:
        return
    kept: list[object] = []
    hidden_prefixes = ("report/", "log/", "cache/")
    for row in rows:
        if not isinstance(row, dict):
            kept.append(row)
            continue
        path = str(row.get("path", row.get("file", row.get("filename", "")))).replace("\\", "/")
        consumers = row.get("consumers")
        if isinstance(consumers, list) and not consumers:
            continue
        if path == report_file or path.startswith(hidden_prefixes) or path in {".run_id", "run_record.json"}:
            continue
        if any(token in path.lower() for token in ("checksum", "provenance", "manifest")):
            continue
        kept.append(row)
    updated = dict(container) if container is not None else {"rows": kept}
    updated[key] = kept
    values["OUTPUT_TABLE_ROWS"] = updated
    values["TABLE:OUTPUTS"] = updated


def _row_values_for_kind(record: object, kind: str, columns: list[str] | None, count: int) -> list[object]:
    keys = columns or list(TABLE_SPECS[kind]["keys"])
    if isinstance(record, dict):
        values = []
        for key in keys[:count]:
            value = record.get(key, "")
            if value in (None, "") and key == "description":
                value = record.get("content", record.get("purpose", ""))
            values.append(value)
        return values + [""] * max(0, count - len(values))
    if isinstance(record, list):
        return list(record[:count]) + [""] * max(0, count - len(record))
    return [_scalar(record)] + [""] * max(0, count - 1)


def _expand_table_rows(document, values: dict[str, object]) -> int:
    """Fill the prototype data row of each named table, preserving its style."""
    from docx.table import Table

    count = 0
    tables: list[Table] = []

    def collect(container):
        for table in getattr(container, "tables", []):
            tables.append(table)
            for row in table.rows:
                for cell in row.cells:
                    collect(cell)

    collect(document)
    for table in tables:
        kind = _table_kind(table)
        if kind is None or len(table.rows) < 2:
            continue
        spec = TABLE_SPECS[kind]
        raw = _lookup(values, spec["rows_key"])
        if raw is None:
            raw = _lookup(values, f"TABLE:{kind}")
        if raw is None:
            continue
        records, columns = _records_and_columns(raw, kind)
        header = table.rows[0]
        if columns and len(columns) == len(header.cells):
            for cell, column in zip(header.cells, columns):
                _cell_set_text(cell, column)
        prototype = table.rows[1]
        template_xml = deepcopy(prototype._tr)
        parent = prototype._tr.getparent()
        insert_at = parent.index(prototype._tr)
        for offset, record in enumerate(records):
            cloned = deepcopy(template_xml)
            cells = cloned.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
            row_values = _row_values_for_kind(record, kind, columns, len(cells))
            for cell, cell_value in zip(cells, row_values):
                text_nodes = cell.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                if text_nodes:
                    text_nodes[0].text = _scalar(cell_value)
                    for node in text_nodes[1:]:
                        node.text = ""
            parent.insert(insert_at + offset, cloned)
        parent.remove(prototype._tr)
        count += len(records)
    return count


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _figure_ids(values: dict[str, object]) -> list[str]:
    found: set[str] = set()
    for key in values:
        match = re.match(r"FIGURE[.:]([A-Z][A-Z0-9_-]*)[.]SOURCE$", key.upper())
        if match and values[key] not in (None, ""):
            found.add(match.group(1))
    return sorted(found, key=lambda item: (int(item[1:]) if item[1:].isdigit() else 10**9, item))


def _replace_xml_text(root, replacements: dict[str, str]) -> None:
    for node in root.iter():
        if node.tag.rsplit("}", 1)[-1] != "t" or node.text is None:
            continue
        text = node.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        node.text = text


def _clone_figure_blocks(document, values: dict[str, object]) -> None:
    """Clone the complete stock F1 title/image/caption block for F2…Fn."""
    ids = _figure_ids(values)
    if len(ids) <= 1:
        return
    paragraphs = list(document.paragraphs)
    source = next((p for p in paragraphs if "[[FIGURE:F1.SOURCE]]" in _paragraph_text(p)), None)
    title = next((p for p in paragraphs if "[[FIGURE:F1.TITLE]]" in _paragraph_text(p)), None)
    caption = next((p for p in paragraphs if "[[FIGURE:F1.CAPTION]]" in _paragraph_text(p)), None)
    if not all((source, title, caption)):
        return
    blocks = [title._p, source._p, caption._p]
    insertion = caption._p
    for number, ident in enumerate(ids[1:], 2):
        clones = [deepcopy(block) for block in blocks]
        replacements = {
            "[[FIGURE:F1.TITLE]]": f"[[FIGURE:{ident}.TITLE]]",
            "[[FIGURE:F1.SOURCE]]": f"[[FIGURE:{ident}.SOURCE]]",
            "[[FIGURE:F1.CAPTION]]": f"[[FIGURE:{ident}.CAPTION]]",
            "图 1.": f"图 {number}.",
        }
        for clone in clones:
            _replace_xml_text(clone, replacements)
            # A cloned block is filled by marker text; duplicate bookmarks
            # would make Word repair the document, so remove them.
            for node in list(clone.iter()):
                if node.tag.rsplit("}", 1)[-1] in {"bookmarkStart", "bookmarkEnd"}:
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)
        for clone in clones:
            insertion.addnext(clone)
            insertion = clone


def _ensure_first_figure(values: dict[str, object]) -> None:
    """Use the first declared figure for the stock F1 block when needed."""
    ids = _figure_ids(values)
    if not ids or "F1" in ids:
        return
    first = ids[0]
    # 模板只有一个原型 F1；若输入使用业务 figure_id，移动（而非复制）
    # 首个图的全部字段，避免后续克隆时把同一张图重复嵌入一次。
    prefixes = (f"FIGURE.{first}.", f"FIGURE:{first}.")
    moved: dict[str, object] = {}
    for key, value in list(values.items()):
        normalised = _key(key)
        if any(normalised.startswith(prefix) for prefix in prefixes):
            suffix = normalised.rsplit(".", 1)[-1]
            moved[f"FIGURE.F1.{suffix}"] = value
            del values[key]
    values.update(moved)


def _ensure_first_note(values: dict[str, object]) -> None:
    """把第一个显式 notes[]/NOTE:<kind> 绑定到模板的方向锚点。"""
    if _slot_text(_lookup(values, "[[NOTE:DIRECTION]]")).strip():
        return
    notes = _explicit_notes(values)
    if notes:
        values["NOTE:DIRECTION"] = notes[0]


def _explicit_notes(values: dict[str, object]) -> list[object]:
    """按输入顺序取得显式 Note，并去掉方向别名造成的重复。"""
    ordered = values.get("_NOTE_ORDER")
    if isinstance(ordered, list):
        return [item for item in ordered if _slot_text(item).strip()]
    found: list[object] = []
    for key, value in values.items():
        name = _key(key)
        if name.startswith("NOTE:") and _slot_text(value).strip():
            if not any(_slot_text(value) == _slot_text(previous) for previous in found):
                found.append(value)
    for key in ("NOTES", "NOTE"):
        value = values.get(key)
        candidates = value if isinstance(value, list) else [value]
        for item in candidates:
            if _slot_text(item).strip() and not any(_slot_text(item) == _slot_text(previous) for previous in found):
                found.append(item)
    return found


def _clone_note_blocks(document, values: dict[str, object]) -> None:
    """为第二个及后续显式 Note 复制完整的独立 callout table。"""
    notes = _explicit_notes(values)
    if len(notes) <= 1:
        return
    note_table = next(
        (table for table in document.tables
         if "[[NOTE:DIRECTION]]" in "".join(cell.text for row in table.rows for cell in row.cells)),
        None,
    )
    if note_table is None:
        return
    template_xml = deepcopy(note_table._element)
    insertion = note_table._element
    for index, note in enumerate(notes[1:], 2):
        ident = "direction"
        if isinstance(note, dict):
            ident = str(note.get("id", f"note{index}")).upper()
        marker = f"[[NOTE:{ident}]]"
        clone = deepcopy(template_xml)
        _replace_xml_text(clone, {"[[NOTE:DIRECTION]]": marker})
        for node in list(clone.iter()):
            if node.tag.rsplit("}", 1)[-1] in {"bookmarkStart", "bookmarkEnd"}:
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)
        insertion.addnext(clone)
        insertion = clone


def _remove_optional_figure_block(document, values: dict[str, object]) -> None:
    """Remove the stock F1 block when no published figure is supplied."""
    source = _lookup(values, "[[FIGURE:F1.SOURCE]]")
    if source not in (None, ""):
        return
    paragraphs = list(_iter_paragraphs(document))
    for paragraph in paragraphs:
        text = _paragraph_text(paragraph)
        if any(token in text for token in ("[[FIGURE:F1.TITLE]]", "[[FIGURE:F1.SOURCE]]", "[[FIGURE:F1.CAPTION]]")):
            _remove_paragraph(paragraph)
    # The heading becomes an empty chapter if the optional figure block is gone.
    for paragraph in list(_iter_paragraphs(document)):
        if _paragraph_text(paragraph) == "图件":
            _remove_paragraph(paragraph)
            break


def _remove_optional_qc(document, values: dict[str, object]) -> None:
    """Remove the QC section when no QC fact was declared."""
    qc = _lookup(values, "[[ANALYSIS_QC]]")
    if qc not in (None, "", [], {}):
        return
    paragraphs = list(document.paragraphs)
    heading = next((p for p in paragraphs if _paragraph_text(p) == "质控与异常"), None)
    if heading is None:
        return
    remove = [heading]
    after_heading = False
    for paragraph in paragraphs:
        if paragraph is heading:
            after_heading = True
            continue
        if not after_heading:
            continue
        if _paragraph_text(paragraph) in {"分析结果与解读", "综合结论"}:
            break
        remove.append(paragraph)
    for paragraph in remove:
        _remove_paragraph(paragraph)


def _remove_optional_note(document, values: dict[str, object]) -> None:
    """Remove the styled Note block when no note is applicable."""
    note = _lookup(values, "[[NOTE:DIRECTION]]")
    if _slot_text(note).strip():
        return
    for table in list(document.tables):
        content = "".join(cell.text for row in table.rows for cell in row.cells)
        first_row = "".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
        if "[[NOTE:" in content or re.search(r"^\s*Note\b", first_row, re.I):
            element = table._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)


def _note_slot_items(values: dict[str, object]):
    """Yield explicit non-empty Note slots without inventing comparison prose."""
    ordered = _explicit_notes(values)
    if ordered:
        for index, value in enumerate(ordered):
            yield f"NOTE[{index}]", _slot_text(value).strip()
        return
    for key, value in values.items():
        name = _key(key)
        if name.startswith("NOTE:"):
            text = _slot_text(value).strip()
            if text:
                yield name, text
        elif name in {"NOTE", "NOTES"}:
            if isinstance(value, list):
                for index, item in enumerate(value):
                    text = _slot_text(item).strip()
                    if text:
                        yield f"{name}[{index}]", text
            else:
                text = _slot_text(value).strip()
                if text:
                    yield name, text


def _validate_render_values(values: dict[str, object], *, root: Path | None, final: bool) -> list[str]:
    """检查直接 slot 输入；pack 模式的 schema 由共享 validator 负责。"""
    if not final:
        return []
    errors: list[str] = []
    required = (
        "REPORT_TITLE", "REPORT_AUDIENCE", "REPORT_SUMMARY", "ANALYSIS_SCOPE",
        "ANALYSIS_METHOD", "ANALYSIS_RESULT", "ANALYSIS_CONCLUSION",
        "ANALYSIS_LIMITATIONS", "OUTPUTS_INTRO", "REFERENCES",
    )
    for name in required:
        value = _lookup(values, f"[[{name}]]")
        if value is None or (name != "ANALYSIS_RESULT" and not _slot_text(value).strip()):
            errors.append(f"required report slot is empty: {name}")
        text = _visible_references(value) if name == "REFERENCES" else _slot_text(value)
        if QUESTION_TEXT.search(text):
            errors.append(f"report slot must be declarative text: {name}")
        if VALUE_PLACEHOLDER.search(text):
            errors.append(f"report slot contains unresolved placeholder: {name}")
        if VISIBLE_SOURCE_URL.search(text):
            errors.append(f"report slot must not expose URL/DOI: {name}")

    for name, text in _note_slot_items(values):
        if GENERIC_NOTE_TEXT.search(text):
            errors.append(f"{name} contains generic Note placeholder; provide a measured fact")
        if QUESTION_TEXT.search(text):
            errors.append(f"{name} must be declarative text")
        if VALUE_PLACEHOLDER.search(text):
            errors.append(f"{name} contains unresolved placeholder")

    # A figure is one complete source/caption unit.  A partial declaration is
    # not silently turned into an empty title or generic caption.
    figure_names = set()
    for key in values:
        match = re.match(r"FIGURE[.:]([A-Z][A-Z0-9_-]*)\.(?:TITLE|SOURCE|CAPTION)$", _key(key))
        if match:
            figure_names.add(match.group(1))
    for ident in sorted(figure_names):
        source = _lookup(values, f"[[FIGURE:{ident}.SOURCE]]")
        title = _lookup(values, f"[[FIGURE:{ident}.TITLE]]")
        caption = _lookup(values, f"[[FIGURE:{ident}.CAPTION]]")
        if source in (None, ""):
            if any(_slot_text(item).strip() for item in (title, caption)):
                errors.append(f"figure {ident} has title/caption without source")
            continue
        if not _slot_text(title).strip():
            errors.append(f"figure {ident} requires title")
        caption_text = _slot_text(caption).strip()
        if not caption_text:
            errors.append(f"figure {ident} requires caption")
        elif QUESTION_TEXT.search(caption_text) or VALUE_PLACEHOLDER.search(caption_text):
            errors.append(f"figure {ident} caption is not declarative evidence")
        if root is not None:
            relative = normalise_relative(_scalar(source))
            if relative is None:
                errors.append(f"figure {ident} source must be a safe relative path")
            else:
                resolved = resolve_under(root, relative)
                if resolved is None or not resolved.is_file():
                    errors.append(f"figure {ident} source file does not exist: {_scalar(source)}")

    for kind, spec in TABLE_SPECS.items():
        raw = _lookup(values, spec["rows_key"])
        if raw is None:
            raw = _lookup(values, f"TABLE:{kind}")
        if raw is None:
            errors.append(f"required table slot is empty: {kind}")
            continue
        records, columns = _records_and_columns(raw, kind)
        count = len(columns) if columns else len(spec["keys"])
        for index, record in enumerate(records):
            row = _row_values_for_kind(record, kind, columns, count)
            if any(not _slot_text(item).strip() for item in row[:count]):
                errors.append(f"{kind} table row {index} contains an empty dynamic field")
    return errors


def _caption_value(values: dict[str, object], marker: str) -> object:
    direct = _lookup(values, marker)
    if direct not in (None, "", [], {}):
        return direct
    body = marker[2:-2].upper()
    match = re.fullmatch(r"TABLE:(RESULTS|OUTPUTS|VERSIONS)\.CAPTION", body)
    if match:
        kind = match.group(1)
        raw = values.get(f"TABLE:{kind}")
        if raw in (None, "", [], {}):
            raw = values.get(TABLE_SPECS[kind]["rows_key"])
        if raw in (None, "", [], {}):
            raw = _lookup(values, f"TABLE:{kind}")
        if isinstance(raw, dict):
            return raw.get("caption", "")
    return direct


def fill_document(template: Path, values: dict[str, object], output: Path, root: Path | None = None) -> dict[str, object]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("python-docx is required for DOCX template rendering") from exc
    document = Document(str(template))
    # 旧版模板可能把 Note 段落装饰误套在已有图件上；先清理所有 drawing
    # 段落，再处理本次 slot，保证旧模板与新模板得到相同的透明图源。
    _clear_all_figure_paragraph_boxes(document)
    _ensure_first_figure(values)
    _ensure_first_note(values)
    _remove_optional_qc(document, values)
    _remove_optional_note(document, values)
    _clone_note_blocks(document, values)
    _clone_figure_blocks(document, values)
    _remove_optional_figure_block(document, values)
    expanded_rows = _expand_table_rows(document, values)
    replaced = 0
    unresolved: list[str] = []
    for paragraph in _iter_paragraphs(document):
        while True:
            full = _paragraph_text(paragraph)
            matches = list(TEMPLATE_MARKER.finditer(full))
            if not matches:
                break
            match = matches[0]
            marker = match.group(0)
            value = _caption_value(values, marker) if marker[2:-2].upper().endswith(".CAPTION") else _lookup(values, marker)
            marker_name = marker[2:-2].upper()
            if value is None and marker_name in {"QC.CONDITION", "RESULTS.CONDITION", "FIGURES.CONDITION", "TABLE:RESULTS", "TABLE:OUTPUTS", "TABLE:VERSIONS"}:
                value = ""
            if marker[2:-2].upper().endswith(SOURCE_SUFFIX) and value not in (None, ""):
                if _insert_picture(paragraph, marker, value, root):
                    replaced += 1
                    continue
                unresolved.append(f"invalid figure source: {value}")
                _replace_span(paragraph, match.start(), match.end(), "")
                replaced += 1
                continue
            if value is None:
                unresolved.append(marker)
                # Leave it visible for draft diagnostics and avoid looping.
                break
            replacement = _visible_references(value) if marker_name == "REFERENCES" else _slot_text(value)
            if marker[2:-2].upper().startswith("NOTE:") and replacement.lower().startswith("note："):
                replacement = replacement.split("：", 1)[1]
            _replace_span(paragraph, match.start(), match.end(), replacement)
            replaced += 1
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", dir=output.parent, prefix=f".{output.name}.", delete=False) as handle:
            temporary = Path(handle.name)
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        if temporary is not None and temporary.exists():
            temporary.unlink()
    return {"replaced": replaced, "expanded_rows": expanded_rows, "unresolved": sorted(set(unresolved))}


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--values", "--slots", dest="values", type=Path)
    parser.add_argument("--evidence-pack", type=Path)
    parser.add_argument("--plan", type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--root", type=Path)
    parser.add_argument("--final", action="store_true", help="run release marker and structure gates")
    parser.add_argument("--require-note", action="store_true")
    parser.add_argument("--json", action="store_true", dest="as_json")
    args = parser.parse_args(argv)
    errors: list[str] = []
    warnings: list[str] = []
    values_document: dict[str, object] | None = None
    pack_for_render: dict[str, object] | None = None
    candidate: Path | None = None
    target = args.output.resolve()
    try:
        template = args.template.resolve()
        if template.is_dir():
            template = template / "report_template.docx"
        if not template.is_file():
            raise ValueError(f"template does not exist: {template}")
        marker_info = extract_template_markers(template)
        if marker_info.get("error"):
            raise ValueError(str(marker_info["error"]))
        bookmark_map = marker_info.get("bookmarks", {})
        if isinstance(bookmark_map, dict):
            for marker in marker_info.get("markers", []):
                marker_body = marker[2:-2].upper()
                # Prototype cell fields are replaced by table-row expansion;
                # they do not need one bookmark per column.  The caption/table
                # anchor is the stable bookmark for the whole table.
                if marker_body.startswith(("RESULT.", "OUTPUT.", "VERSION.")):
                    continue
                expected = _marker_bookmark_name(marker)
                actual = bookmark_map.get(marker, [])
                if expected not in actual:
                    message = f"template marker {marker} has no matching bookmark {expected}"
                    (errors if args.final else warnings).append(message)
        if args.values:
            loaded_values = _load(args.values)
            if not isinstance(loaded_values, dict):
                raise ValueError("slot values must be a JSON object")
            values_document = loaded_values
            values = {}
            if args.evidence_pack:
                pack_for_defaults = _load(args.evidence_pack)
                if not isinstance(pack_for_defaults, dict):
                    raise ValueError("evidence pack must be a JSON object")
                pack_for_render = pack_for_defaults
                values.update(values_from_pack(pack_for_defaults))
            values.update(_flatten_values(loaded_values))
        elif args.evidence_pack:
            pack = _load(args.evidence_pack)
            if not isinstance(pack, dict):
                raise ValueError("evidence pack must be a JSON object")
            pack_for_render = pack
            values = values_from_pack(pack)
        else:
            raise ValueError("one of --values/--evidence-pack is required")
        report_file = None
        if args.plan:
            plan_for_output = _load(args.plan)
            if isinstance(plan_for_output, dict):
                policy = plan_for_output.get("output_policy")
                if isinstance(policy, dict) and isinstance(policy.get("report_file"), str):
                    report_file = normalise_relative(policy["report_file"])
        _filter_public_output_rows(values, report_file)
        root = args.root.resolve() if args.root else None
        if root is not None and not root.is_dir():
            raise ValueError(f"artifact root does not exist: {root}")
        # A template with slot bookmarks is the DOCX-first contract.  Legacy
        # templates without bookmarks remain usable but are reported by the
        # structural validator in release mode.
        # Final values-only runs still need an explicit release declaration;
        # otherwise a caller could turn a draft slot map into PASS by adding
        # ``--final`` without any evidence gate.
        if args.final:
            declared_profile = values_document.get("quality_profile") if values_document else None
            if values_document is not None and args.evidence_pack is None and declared_profile != "release":
                errors.append("final template render requires slot values quality_profile=release")
            if values_document is not None and args.evidence_pack is not None and declared_profile not in (None, "release"):
                errors.append("slot values quality_profile must be release when supplied")
            if pack_for_render is not None and pack_for_render.get("quality_profile") != "release":
                errors.append("final template render requires evidence pack quality_profile=release")
            errors.extend(_validate_render_values(values, root=root, final=True))
        # release 先写同目录候选文件；所有门禁通过后才原子替换正式报告。
        # 失败时保留已有正式报告，不把半成品当作新交付。
        render_path = target
        if args.final:
            target.parent.mkdir(parents=True, exist_ok=True)
            with tempfile.NamedTemporaryFile(
                suffix=target.suffix or ".docx",
                prefix=f".{target.name}.",
                dir=target.parent,
                delete=False,
            ) as handle:
                candidate = Path(handle.name)
            candidate.unlink()
            render_path = candidate
        result = fill_document(template, values, render_path, root)
        checked = inspect(render_path, final=args.final, require_note=args.require_note)
        result["figure_box_issues"] = checked.get("figure_box_issues", [])
        if args.final:
            if args.evidence_pack:
                pack = pack_for_render if pack_for_render is not None else _load(args.evidence_pack)
                if not isinstance(pack, dict):
                    raise ValueError("evidence pack must be a JSON object")
                errors.extend(validate_pack(pack, root, require_files=True, final=True))
                if args.plan:
                    plan = _load(args.plan)
                    if not isinstance(plan, dict):
                        raise ValueError("report plan must be a JSON object")
                    point_ids = {item.get("id") for item in pack.get("analysis_points", []) if isinstance(item, dict)}
                    errors.extend(validate_plan(plan, point_ids, final=True))
                    errors.extend(validate_slot_alignment(pack, plan, strict=True))
            if result["unresolved"]:
                errors.append(f"unresolved template markers: {result['unresolved']}")
            checked = inspect(render_path, final=True, require_note=args.require_note)
            if checked.get("status") != "PASS":
                errors.extend(f"DOCX: {message}" for message in checked.get("errors", []))
                if not checked.get("errors") and checked.get("status") != "PASS":
                    errors.append(f"DOCX structural gate returned {checked.get('status')}")
        status = "PASS" if args.final and not errors else ("BLOCKED" if errors else "DRAFT")
        if args.final and status == "PASS":
            os.replace(render_path, target)
            candidate = None
        elif args.final and candidate is not None and candidate.exists():
            candidate.unlink()
            candidate = None
        payload = {"status": status, **result, "template_markers": marker_info.get("markers", []), "warnings": warnings, "errors": errors}
        if args.as_json:
            print(json.dumps(payload, ensure_ascii=False, indent=2))
        else:
            output_for_message = target if args.final else args.output.resolve()
            digest = hashlib.sha256(output_for_message.read_bytes()).hexdigest() if output_for_message.is_file() else ""
            print(f"DOCX_TEMPLATE_{status} output={output_for_message} sha256={digest}")
            for message in errors:
                print(message, file=sys.stderr)
            for message in warnings:
                print(message)
        return 0 if status == "PASS" else 2
    except (OSError, ValueError, json.JSONDecodeError, RuntimeError) as exc:
        if candidate is not None and candidate.exists():
            candidate.unlink()
        print(f"DOCX_TEMPLATE_BLOCKED: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
