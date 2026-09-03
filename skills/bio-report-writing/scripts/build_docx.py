#!/usr/bin/env python3
"""Build a conservative draft preview from a report plan and evidence pack.

This generic key-value builder is draft-only. Formal reports come from the
module's R/Python renderer. It never performs scientific calculation, infers
missing values, or crops source figures.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
from pathlib import Path
import sys
import tempfile

from render_text import (
    compact,
    comparison_text,
    method_text,
    parameters_text,
    reference_texts,
    result_texts,
)
from validate_report_contract import (
    normalise_relative,
    resolve_under,
    section_semantics,
    validate_flat_result_paths,
    validate_pack,
    validate_plan,
)


def read(path: Path) -> dict:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return value


def value_text(value: object) -> str:
    if value is None or value == "" or value == [] or value == {}:
        return "[EVIDENCE_REQUIRED]"
    return compact(value)


def set_font(style, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def set_keep(paragraph, *, next_line: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = next_line
    paragraph.paragraph_format.keep_together = together


def set_row_controls(row, repeat_header: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)
    if repeat_header:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)


def add_label_paragraph(document, label: str, content: object) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}：")
    run.bold = True
    paragraph.add_run(value_text(content))
    set_keep(paragraph, together=True)


def inside(root: Path, relative: str) -> Path:
    normalised = normalise_relative(relative)
    path = resolve_under(root, normalised) if normalised else None
    if path is None:
        raise ValueError(f"path must be a safe relative path under evidence root: {relative}")
    return path


def add_figure(document, reference: dict, root: Path) -> None:
    path_value = reference.get("path")
    paragraph = document.add_paragraph()
    if not isinstance(path_value, str) or not path_value:
        paragraph.add_run("[EVIDENCE_REQUIRED:figure_path]")
        return
    path = inside(root, path_value)
    if not path.is_file() or path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
        paragraph.add_run(f"[EVIDENCE_REQUIRED:figure {reference.get('id', '')}]")
        return
    try:
        from PIL import Image

        with Image.open(path) as image:
            width_px, height_px = image.size
    except Exception as exc:  # pragma: no cover - depends on image codecs
        paragraph.add_run(f"[FIGURE_READ_FAILED:{path.name}:{exc}]")
        return
    max_width = 6.5
    max_height = 8.1
    width = max_width
    height = width * height_px / width_px if width_px else max_height
    if height > max_height:
        height = max_height
        width = height * width_px / height_px
    from docx.shared import Inches

    paragraph.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    set_keep(paragraph, next_line=True, together=True)
    caption = document.add_paragraph()
    caption.style = "Caption"
    caption.add_run(f"{reference.get('id', 'Figure')}: {value_text(reference.get('caption'))}")
    set_keep(caption, together=True)


def add_figure_or_file(document, reference: dict, root: Path) -> None:
    """Embed an image, or describe a non-image table/file without a fake image."""
    path_value = reference.get("path") if isinstance(reference, dict) else None
    if isinstance(path_value, str) and Path(path_value).suffix.lower() in {".png", ".jpg", ".jpeg"}:
        add_figure(document, reference, root)
        return
    paragraph = document.add_paragraph()
    ident = value_text(reference.get("id", "Figure/Table"))
    path = value_text(path_value)
    paragraph.add_run(f"{ident}：{path}")
    set_keep(paragraph, together=True)


def configure(document) -> None:
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm

    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    styles = document.styles
    set_font(styles["Normal"], "Times New Roman", 11.5)
    set_font(styles["Title"], "黑体", 22, True, "0F4761")
    set_font(styles["Heading 1"], "黑体", 16, True, "0F4761")
    set_font(styles["Heading 2"], "黑体", 14, True, "0F4761")
    set_font(styles["Heading 3"], "黑体", 12, True, "0F4761")
    set_font(styles["Caption"], "Times New Roman", 9.5, True, "2F75B5")


def build(plan: dict, pack: dict, output: Path, root: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX building") from exc
    document = Document()
    configure(document)
    title = document.add_paragraph(style="Title")
    title.alignment = 1
    title.add_run(value_text(plan.get("title")))
    banner = document.add_paragraph("DRAFT：通用 slot 预览；正式交付请使用模块 renderer。")
    banner.alignment = 1
    banner.runs[0].italic = True
    points = {point.get("id"): point for point in pack.get("analysis_points", []) if isinstance(point, dict)}
    for section in plan.get("sections", []):
        sid = section.get("id")
        heading = document.add_heading(value_text(section.get("title")), level=1)
        set_keep(heading, next_line=True)
        ids = section.get("analysis_point_ids", [])
        semantics = section_semantics(section)
        if sid == "summary" or "summary" in semantics:
            for point in points.values():
                add_label_paragraph(
                    document,
                    str(point.get("title", point.get("id", "分析点"))),
                    result_texts(point.get("results"), point.get("id", "AP") + ".results"),
                )
        elif sid == "scope" or "scope" in semantics:
            for point in points.values():
                add_label_paragraph(document, str(point.get("title", point.get("id", "分析点"))), point.get("scope"))
        elif sid in {"methods", "qc", "results", "conclusion", "limitations"} or semantics.intersection({"method", "qc", "results", "interpretation", "limitations"}):
            if "method" in semantics:
                sid = "methods"
            elif "qc" in semantics:
                sid = "qc"
            elif "limitations" in semantics:
                sid = "limitations"
            elif "interpretation" in semantics and "results" not in semantics:
                sid = "conclusion"
            else:
                sid = "results"
            for point_id in ids:
                point = points.get(point_id)
                if point is None:
                    add_label_paragraph(document, "证据目标", f"[EVIDENCE_REQUIRED:{point_id}]")
                    continue
                sub = document.add_heading(value_text(point.get("title")), level=2)
                set_keep(sub, next_line=True)
                if sid == "methods":
                    add_label_paragraph(document, "输入与范围", point.get("scope"))
                    add_label_paragraph(document, "方法与版本", method_text(point.get("method"), point.get("id", "AP") + ".method"))
                    add_label_paragraph(document, "参数与统计口径", parameters_text(point.get("parameters"), point.get("id", "AP") + ".parameters"))
                    add_label_paragraph(document, "推断单位与比较方向", comparison_text(point, point.get("id", "AP") + ".comparison"))
                elif sid == "results":
                    add_label_paragraph(document, "结果", result_texts(point.get("results"), point.get("id", "AP") + ".results"))
                    for reference in point.get("figure_table_refs", []):
                        if isinstance(reference, dict):
                            add_figure_or_file(document, reference, root)
                    add_label_paragraph(document, "领域解释", point.get("interpretation"))
                    add_label_paragraph(document, "下一步用途", point.get("next_step"))
                elif sid == "qc":
                    add_label_paragraph(document, "质控与异常", point.get("qc"))
                elif sid == "conclusion":
                    add_label_paragraph(document, "结论回收", point.get("interpretation"))
                else:
                    add_label_paragraph(document, "解释边界", point.get("limitations"))
                    add_label_paragraph(document, "待验证用途", point.get("next_step"))
        elif sid == "outputs" or "outputs" in semantics:
            outputs = [item for point in points.values() for item in point.get("outputs", []) if isinstance(item, dict) and item.get("published")]
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            set_row_controls(table.rows[0], repeat_header=True)
            for cell, label in zip(table.rows[0].cells, ("文件", "类型", "用途")):
                cell.text = label
            for item in outputs:
                if normalise_relative(item.get("path")) == normalise_relative(plan.get("output_policy", {}).get("report_file")):
                    continue
                cells = table.add_row().cells
                cells[0].text = value_text(item.get("path"))
                cells[1].text = value_text(item.get("kind"))
                cells[2].text = value_text(item.get("purpose"))
                set_row_controls(table.rows[-1])
        elif sid == "references" or "references" in semantics:
            add_label_paragraph(document, "实际使用来源", reference_texts(pack.get("references"), "references"))
        elif sid == "versions" or "versions" in semantics:
            add_label_paragraph(document, "软件与资源版本", reference_texts(pack.get("versions"), "versions"))
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", dir=output.parent, prefix=f".{output.name}.", delete=False) as handle:
            temporary = Path(handle.name)
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        if temporary is not None and temporary.exists():
            temporary.unlink()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", required=True, type=Path)
    parser.add_argument("--evidence-pack", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--root", type=Path, help="root for relative figure paths; defaults to evidence pack directory")
    parser.add_argument("--final", action="store_true", help="require release/evidence-complete inputs before building")
    args = parser.parse_args(argv)
    try:
        plan = read(args.plan)
        pack = read(args.evidence_pack)
        root = (args.root or args.evidence_pack.parent).resolve()
        if not root.is_dir():
            raise ValueError(f"artifact root does not exist or is not a directory: {root}")
        point_ids = {item.get("id") for item in pack.get("analysis_points", []) if isinstance(item, dict)}
        preflight = validate_pack(pack, root, require_files=args.final) + validate_plan(plan, point_ids)
        if pack.get("result_layout") != plan.get("result_layout"):
            preflight.append("plan and evidence pack result_layout must match")
        if plan.get("result_layout") == "flat":
            preflight += validate_flat_result_paths(pack)
        pack_ref = plan.get("evidence_pack") if isinstance(plan, dict) else None
        if isinstance(pack_ref, str) and pack_ref.strip() and (args.plan.parent / pack_ref.replace("\\", "/")).resolve() != args.evidence_pack.resolve():
            preflight.append("plan evidence_pack does not match --evidence-pack")
        if args.final:
            preflight.append("generic key-value builder is draft-only; invoke the module R/Python renderer for final output")
        if args.final and isinstance(pack, dict):
            if pack.get("quality_profile") != "release" or plan.get("quality_profile") != "release":
                preflight.append("final build requires quality_profile=release")
            for index, point in enumerate(pack.get("analysis_points", [])):
                if isinstance(point, dict) and (point.get("status") not in {"complete", "valid_no_findings"} or not point.get("limitations") or not isinstance(point.get("next_step"), str) or not point["next_step"].strip()):
                    preflight.append(f"final build requires complete analysis point {index}")
        if preflight:
            raise ValueError("preflight: " + "; ".join(preflight))
        build(plan, pack, args.output, root)
    except (OSError, ValueError, json.JSONDecodeError, RuntimeError) as exc:
        print(f"DOCX_BUILD_BLOCKED: {exc}", file=sys.stderr)
        return 2
    data = args.output.read_bytes()
    print(f"DOCX_BUILD_DRAFT output={args.output.resolve()} bytes={len(data)} sha256={hashlib.sha256(data).hexdigest()}")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
