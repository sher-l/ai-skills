#!/usr/bin/env python3
"""Build the portable DOCX-first report template and style reference.

``report_template.docx`` is intentionally boring: a renderer can locate every
slot by its visible ``[[...]]`` marker or the accompanying Word bookmark,
replace the marker, and keep the fixed reader-facing chapter order.  The
sky-blue Note is emitted as explicit WordprocessingML so its fill, border and
label colour remain inspectable after a Quarto/Pandoc render.
"""
from __future__ import annotations

from pathlib import Path
import sys


NOTE_BORDER = "5B9BD5"
NOTE_FILL = "DDEBF7"
NOTE_LABEL = "2F75B5"
HEADING = "0F4761"


def _imports():
    """Load the optional DOCX dependency only when this generator runs."""
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise RuntimeError("install optional python-docx") from exc
    return {
        "Document": Document,
        "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Cm": Cm,
        "Pt": Pt,
        "RGBColor": RGBColor,
    }


def _set(element, qn, name: str, value: str) -> None:
    element.set(qn(name if ":" in name else f"w:{name}"), value)


def _run_font(
    run,
    api,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    name: str | None = None,
) -> None:
    # 未显式指定时继承段落样式，避免 marker 覆盖标题、正文和图注字号。
    if name is not None:
        run.font.name = name
    if size is not None:
        run.font.size = api["Pt"](size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = api["RGBColor"].from_string(color)
    if name is not None:
        run._element.get_or_add_rPr().rFonts.set(api["qn"]("w:eastAsia"), "宋体")


def _configure_styles(document, api) -> None:
    qn, Pt, RGBColor = api["qn"], api["Pt"], api["RGBColor"]
    for name, size, bold, color in (
        ("Normal", 11.5, False, None),
        ("Title", 22, True, HEADING),
        ("Heading 1", 16, True, HEADING),
        ("Heading 2", 14, True, HEADING),
        ("Heading 3", 12, True, HEADING),
        # 图题/图注不是 Note；保持普通深色文字，天蓝色只属于 callout。
        ("Caption", 9.5, False, None),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        if name == "Caption":
            style.font.italic = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if color:
            style.font.color.rgb = RGBColor.from_string(color)


def _configure_page(document, api) -> None:
    section = document.sections[0]
    section.page_width = api["Cm"](21.0)
    section.page_height = api["Cm"](29.7)
    section.top_margin = api["Cm"](2.2)
    section.bottom_margin = api["Cm"](2.0)
    section.left_margin = api["Cm"](2.0)
    section.right_margin = api["Cm"](2.0)


def _bookmark(
    paragraph,
    name: str,
    text: str,
    bookmark_id: int,
    api,
    *,
    color: str | None = None,
    bold: bool | None = None,
    size: float | None = None,
):
    """Append a visible slot marker wrapped in a stable Word bookmark."""
    qn, OxmlElement = api["qn"], api["OxmlElement"]
    start = OxmlElement("w:bookmarkStart")
    _set(start, qn, "id", str(bookmark_id))
    _set(start, qn, "name", name)
    end = OxmlElement("w:bookmarkEnd")
    _set(end, qn, "id", str(bookmark_id))
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    _run_font(run, api, size=size, color=color, bold=bold)
    paragraph._p.append(end)
    return run


class _Bookmarks:
    def __init__(self, api):
        self.api = api
        self.next_id = 1

    def add(self, paragraph, name: str, text: str, **kwargs):
        run = _bookmark(paragraph, name, text, self.next_id, self.api, **kwargs)
        self.next_id += 1
        return run


def _add_heading(document, text: str, level: int):
    return document.add_heading(text, level=level)


def _slot(document, bookmarks, marker: str, bookmark_name: str, *, style: str | None = None, align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    bookmarks.add(paragraph, bookmark_name, marker)
    return paragraph


def _clear_figure_paragraph_box(paragraph, api) -> None:
    """图件段落保持透明、无边框和无额外缩进；图片只按版心等比缩放。"""
    qn = api["qn"]
    ppr = paragraph._p.get_or_add_pPr()
    for name in ("pBdr", "shd", "ind"):
        node = ppr.find(qn(f"w:{name}"))
        if node is not None:
            ppr.remove(node)


def _cell_box(
    cell,
    api,
    *,
    color: str = NOTE_BORDER,
    fill: str = NOTE_FILL,
    left_size: str = "8",
) -> None:
    """Write exact fill and four borders on a table cell."""
    qn, OxmlElement = api["qn"], api["OxmlElement"]
    tcpr = cell._tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tcpr.append(shading)
    _set(shading, qn, "val", "clear")
    _set(shading, qn, "fill", fill)
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        size = left_size if edge == "left" else "8"
        for key, value in (("val", "single"), ("sz", size), ("space", "0"), ("color", color)):
            _set(node, qn, key, value)
    parent = cell._tc.getparent()
    trpr = parent.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def _note_table_borders(table, api, *, color: str = NOTE_BORDER, left_size: str = "16") -> None:
    """在 Note table 层声明外框；insideH/insideV 不画分隔线，接近 DEGs callout。"""
    qn, OxmlElement = api["qn"], api["OxmlElement"]
    # python-docx 的 CT_Tbl 没有 get_or_add_tblPr；新表总会带 tblPr。
    tblpr = table._tbl.tblPr
    if tblpr is None:  # pragma: no cover - 防御自定义旧模板代理
        tblpr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblpr)
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        size = left_size if edge == "left" else "8"
        for key, value in (("val", "single"), ("sz", size), ("space", "0"), ("color", color)):
            _set(node, qn, key, value)
    for edge in ("insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        _set(node, qn, "val", "nil")


def _cell_margins(cell, api, value: str = "120") -> None:
    qn, OxmlElement = api["qn"], api["OxmlElement"]
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        _set(node, qn, "w", value)
        _set(node, qn, "type", "dxa")


def _row_rules(row, api, *, header: bool = False) -> None:
    qn, OxmlElement = api["qn"], api["OxmlElement"]
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))
    if header and trpr.find(qn("w:tblHeader")) is None:
        trpr.append(OxmlElement("w:tblHeader"))


def _keep_next(paragraph, api, value: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    node = ppr.find(api["qn"]("w:keepNext"))
    if value and node is None:
        ppr.append(api["OxmlElement"]("w:keepNext"))
    elif not value and node is not None:
        ppr.remove(node)


def _style_table(table, api, *, header=True, alignments=None) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for index, row in enumerate(table.rows):
        _row_rules(row, api, header=header and index == 0)
        for column, cell in enumerate(row.cells):
            _cell_margins(cell, api, "90")
            cell.vertical_alignment = api["WD_CELL_VERTICAL_ALIGNMENT"].CENTER
            for paragraph in cell.paragraphs:
                if alignments and index > 0:
                    paragraph.alignment = alignments[min(column, len(alignments) - 1)]
                for run in paragraph.runs:
                    _run_font(
                        run,
                        api,
                        size=10.5,
                        bold=index == 0,
                        color=HEADING if index == 0 else None,
                        name="Times New Roman",
                    )


def _set_table_widths(table, api, widths_cm: tuple[float, ...]) -> None:
    """Keep long output headers readable in the fixed A4 text width."""
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = api["Cm"](width)


def _add_note(document, bookmarks, api):
    """添加独立的两行 Note callout；颜色只写入 callout table 的单元格。"""
    table = document.add_table(rows=2, cols=1)
    table.alignment = api["WD_TABLE_ALIGNMENT"].CENTER
    table.autofit = False
    _note_table_borders(table, api, left_size="16")
    for row in table.rows:
        _row_rules(row, api)
        cell = row.cells[0]
        cell.vertical_alignment = api["WD_CELL_VERTICAL_ALIGNMENT"].CENTER
        _cell_box(cell, api, left_size="16")
        _cell_margins(cell, api)
    label_paragraph = table.cell(0, 0).paragraphs[0]
    # 使用可移植文字标签；不依赖外部图标字体，避免目标阅读器出现 tofu。
    label = label_paragraph.add_run("Note：")
    _run_font(label, api, color=NOTE_LABEL, bold=True, name="Times New Roman")
    body_paragraph = table.cell(1, 0).paragraphs[0]
    bookmarks.add(body_paragraph, "slot_note_direction", "[[NOTE:DIRECTION]]")
    return table


def _add_caption(document, bookmarks, figure_id: str = "F1"):
    paragraph = document.add_paragraph(style="Caption")
    bookmarks.add(paragraph, f"slot_figure_{figure_id.lower()}_caption", f"[[FIGURE:{figure_id}.CAPTION]]")
    return paragraph


def _add_figure_block(document, bookmarks, api, figure_id: str = "F1", number: int = 1):
    heading = document.add_heading(level=2)
    heading.add_run(f"图 {number}. ")
    bookmarks.add(heading, f"slot_figure_{figure_id.lower()}_title", f"[[FIGURE:{figure_id}.TITLE]]")
    source = document.add_paragraph()
    source.alignment = api["WD_ALIGN_PARAGRAPH"].CENTER
    # 该标记由完整源图替换；有更多图件时，renderer 复制整个三段图块，
    # 没有图件时删除整块，不保留空章节或占位图。
    bookmarks.add(source, f"slot_figure_{figure_id.lower()}_source", f"[[FIGURE:{figure_id}.SOURCE]]")
    _clear_figure_paragraph_box(source, api)
    _keep_next(source, api)
    caption = _add_caption(document, bookmarks, figure_id)
    _keep_next(heading, api)
    _keep_next(caption, api)


def _add_result_table(document, bookmarks, api):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run("表 1. ")
    bookmarks.add(paragraph, "slot_table_results_caption", "[[TABLE:RESULTS.CAPTION]]")
    table = document.add_table(rows=2, cols=4)
    headers = ("指标", "数值", "单位", "结果来源")
    values = ("[[RESULT.NAME]]", "[[RESULT.VALUE]]", "[[RESULT.UNIT]]", "[[RESULT.SOURCE]]")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value
    align = (api["WD_ALIGN_PARAGRAPH"].LEFT, api["WD_ALIGN_PARAGRAPH"].RIGHT,
             api["WD_ALIGN_PARAGRAPH"].LEFT, api["WD_ALIGN_PARAGRAPH"].LEFT)
    _style_table(table, api, alignments=align)
    _set_table_widths(table, api, (4.2, 2.5, 2.5, 7.8))
    return table


def _add_output_table(document, bookmarks, api):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run("表 2. ")
    bookmarks.add(paragraph, "slot_table_outputs_caption", "[[TABLE:OUTPUTS.CAPTION]]")
    table = document.add_table(rows=2, cols=5)
    headers = ("文件名", "类型", "内容", "用途", "消费者")
    values = ("[[OUTPUT.PATH]]", "[[OUTPUT.KIND]]", "[[OUTPUT.DESCRIPTION]]", "[[OUTPUT.PURPOSE]]", "[[OUTPUT.CONSUMERS]]")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value
    align = tuple(api["WD_ALIGN_PARAGRAPH"].LEFT for _ in range(5))
    _style_table(table, api, alignments=align)
    _set_table_widths(table, api, (3.4, 1.5, 3.7, 4.0, 4.4))
    return table


def _add_version_table(document, bookmarks, api):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run("表 3. ")
    bookmarks.add(paragraph, "slot_table_versions_caption", "[[TABLE:VERSIONS.CAPTION]]")
    table = document.add_table(rows=2, cols=3)
    headers = ("软件或资源", "版本", "用途或来源")
    values = ("[[VERSION.NAME]]", "[[VERSION.VALUE]]", "[[VERSION.PURPOSE]]")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value
    align = tuple(api["WD_ALIGN_PARAGRAPH"].LEFT for _ in range(3))
    _style_table(table, api, alignments=align)
    _set_table_widths(table, api, (5.2, 3.0, 8.8))
    return table


def build_report_template(api):
    """Create the full reusable DOCX template."""
    document = api["Document"]()
    _configure_page(document, api)
    _configure_styles(document, api)
    bookmarks = _Bookmarks(api)

    title = document.add_paragraph(style="Title")
    title.alignment = api["WD_ALIGN_PARAGRAPH"].CENTER
    bookmarks.add(title, "slot_report_title", "[[REPORT_TITLE]]")
    metadata = document.add_paragraph()
    metadata.add_run("读者：").bold = True
    bookmarks.add(metadata, "slot_report_audience", "[[REPORT_AUDIENCE]]")

    for heading, prose, marker, name in (
        ("摘要", "本报告概述本次分析的对象、方法、主要结果和适用范围。", "[[REPORT_SUMMARY]]", "slot_report_summary"),
        ("数据范围与分析边界", "纳入数据、推断单位、比较方向和适用边界如下。", "[[ANALYSIS_SCOPE]]", "slot_analysis_scope"),
        ("材料与方法", "本次分析使用已批准的输入、方法和参数。", "[[ANALYSIS_METHOD]]", "slot_analysis_method"),
        ("质控与异常", "以下列出可核对的质量控制事实及其对解释的影响。", "[[ANALYSIS_QC]]", "slot_analysis_qc"),
        ("分析结果与解读", "结果按分析点的公开顺序呈现，先给结论，再给数字证据和解释边界。", "[[ANALYSIS_RESULT]]", "slot_analysis_result"),
    ):
        _add_heading(document, heading, 1)
        paragraph = document.add_paragraph()
        paragraph.add_run(prose + " ")
        bookmarks.add(paragraph, name, marker)
        if heading == "分析结果与解读":
            _add_note(document, bookmarks, api)
            _add_result_table(document, bookmarks, api)
            # Figure 是结果内容单元，置于综合结论之前；无真实图件时由
            # renderer 删除整个图块，不留下空标题或占位图。
            _add_heading(document, "图件", 2)
            _add_figure_block(document, bookmarks, api, "F1", 1)

    _add_heading(document, "综合结论", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("综合结论仅回收上文已确认的结果。 ")
    bookmarks.add(paragraph, "slot_analysis_conclusion", "[[ANALYSIS_CONCLUSION]]")

    _add_heading(document, "局限、未完成与待验证", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("下列限制说明结果仍可解释的范围和待验证事项。 ")
    bookmarks.add(paragraph, "slot_analysis_limitations", "[[ANALYSIS_LIMITATIONS]]")

    _add_heading(document, "输出文件说明", 1)
    _slot(document, bookmarks, "[[OUTPUTS_INTRO]]", "slot_outputs_intro")
    _add_output_table(document, bookmarks, api)

    _add_heading(document, "参考文献", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("本报告实际使用的来源如下。 ")
    bookmarks.add(paragraph, "slot_references", "[[REFERENCES]]")

    _add_heading(document, "软件与资源版本", 1)
    _add_version_table(document, bookmarks, api)
    return document


def build_style_reference(api):
    """Create a compact style sample retained for old one-off workflows."""
    document = api["Document"]()
    _configure_page(document, api)
    _configure_styles(document, api)
    bookmarks = _Bookmarks(api)
    title = document.add_paragraph(style="Title")
    title.alignment = api["WD_ALIGN_PARAGRAPH"].CENTER
    title.add_run("生信分析报告样式参考")
    _add_heading(document, "分析结果与解读", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("结果：").bold = True
    paragraph.add_run("此处仅放来自 evidence pack 的已确认事实。")
    _add_note(document, bookmarks, api)
    caption = document.add_paragraph(style="Caption")
    caption.add_run("图 1. 完整源图及其逐 panel 图注。")
    _add_heading(document, "软件与资源版本", 1)
    table = document.add_table(rows=2, cols=3)
    for cell, value in zip(table.rows[0].cells, ("软件/资源", "版本", "用途")):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, ("[[VERSION.NAME]]", "[[VERSION.VALUE]]", "[[VERSION.PURPOSE]]")):
        cell.text = value
    _style_table(table, api)
    return document


def main() -> int:
    try:
        api = _imports()
    except RuntimeError as exc:
        print(f"STYLE_REFERENCE_BLOCKED: {exc}", file=sys.stderr)
        return 2
    assets = Path(__file__).resolve().parents[1] / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    template_path = assets / "report_template.docx"
    style_path = assets / "report-style-reference.docx"
    build_report_template(api).save(template_path)
    build_style_reference(api).save(style_path)
    print(f"REPORT_TEMPLATE_PASS output={template_path}")
    print(f"STYLE_REFERENCE_PASS output={style_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
