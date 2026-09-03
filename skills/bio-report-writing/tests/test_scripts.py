#!/usr/bin/env python3
from __future__ import annotations

import json
import base64
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile
import unittest
import zipfile


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"


def run(script: str, *args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run([sys.executable, str(SCRIPTS / script), *args], text=True, capture_output=True, check=False)


def pack(path: Path) -> None:
    value = {
        "schema_version": "0.1.0",
        "module": "demo",
        "quality_profile": "draft",
        "result_layout": "flat",
        "evidence_targets": [{"id": "ET-01", "title": "主要分析结果"}],
        "analysis_points": [{
            "id": "AP-01", "title": "差异表达结果", "scope": "两个供体组的表达矩阵", "qc": "输入和顺序已核对",
            "inputs": [{"id": "IN-01", "path": "result/input.tsv", "identity": "demo"}],
            "method": {"name": "limma", "version": "3.58"}, "parameters": {"fdr": 0.05},
            "statistical_unit": "sample", "comparison": {"target": "case", "reference": "control", "direction": "case-control"},
            "results": [{"name": "n", "value": 4, "unit": "gene", "source": "result/01.DEG_all.tsv"}],
            "outputs": [{"id": "OUT-01", "path": "result/01.DEG_all.tsv", "kind": "table", "published": True, "purpose": "结果表", "consumers": ["project_analyst"]}],
            "figure_table_refs": [], "interpretation_level": "descriptive", "interpretation": "当前数据中的描述性结果", "limitations": ["样本范围"], "status": "complete", "next_step": "作为后续通路分析的输入"
        }]
    }
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


class ReportSkillSmoke(unittest.TestCase):
    def test_unified_cli_exposes_fixed_operations(self) -> None:
        result = run("bio_report.py", "--help")
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("skeleton", result.stdout)

    def test_fixed_pipeline_builds_draft(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            output = root / "report-work"
            pack(evidence)
            result = run("bio_report.py", "run", "--evidence-pack", str(evidence), "--output-dir", str(output))
            self.assertNotEqual(result.returncode, 0, result.stderr)
            self.assertIn("EVIDENCE_NEEDED", result.stdout)
            self.assertTrue((output / "report_plan.json").is_file())
            self.assertTrue((output / "report_draft.md").is_file())
            self.assertTrue((output / "report.docx").is_file())

    def test_plan_skeleton_validate_and_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            markdown = root / "draft.md"
            docx = root / "report.docx"
            pack(evidence)
            result = run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            self.assertEqual(result.returncode, 0, result.stderr)
            result = run("build_report_skeleton.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--output", str(markdown))
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("## 摘要", markdown.read_text(encoding="utf-8"))
            self.assertIn("n=4 gene", markdown.read_text(encoding="utf-8"))
            result = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--markdown", str(markdown))
            self.assertNotEqual(result.returncode, 0)  # skeleton intentionally has evidence markers
            result = run("build_docx.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--output", str(docx))
            self.assertNotEqual(result.returncode, 0, result.stderr)
            self.assertIn("DRAFT", result.stdout)
            result = run("validate_docx_structure.py", str(docx))
            self.assertNotEqual(result.returncode, 0, result.stderr)
            self.assertIn("EVIDENCE_NEEDED", result.stdout)

    def test_question_and_placeholder_are_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            markdown = root / "draft.md"
            pack(evidence)
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            markdown.write_text("# 哪些结果？\n\n[[EVIDENCE_REQUIRED:x]]\n", encoding="utf-8")
            result = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--markdown", str(markdown))
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("declarative", result.stderr)

    def test_docx_visible_question_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            docx = root / "report.docx"
            pack(evidence)
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            value = json.loads(plan.read_text(encoding="utf-8"))
            value["title"] = "哪些结果？"
            plan.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            built = run("build_docx.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--output", str(docx))
            self.assertNotEqual(built.returncode, 0)
            self.assertIn("declarative", built.stderr)

    def test_docx_question_is_blocked_without_final_flag(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            docx = Path(temp) / "question.docx"
            xml = """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>当前 GRN 优先哪些 TF-target 关系？</w:t></w:r></w:p></w:body></w:document>"""
            with zipfile.ZipFile(docx, "w") as archive:
                archive.writestr("word/document.xml", xml)
            checked = run("validate_docx_structure.py", str(docx))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("non-declarative", checked.stderr)

    def test_engineering_status_does_not_enter_reader_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            docx = Path(temp) / "status.docx"
            from docx import Document

            document = Document()
            document.add_paragraph("方法状态：complete；通过运行记录校验。")
            document.save(docx)
            checked = run("validate_docx_structure.py", str(docx), "--final", "--json")
            self.assertNotEqual(checked.returncode, 0)
            payload = json.loads(checked.stdout)
            self.assertGreater(payload["engineering_text"], 0)

    def test_failed_skeleton_keeps_previous_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            output = root / "report.md"
            output.write_text("trusted previous report\n", encoding="utf-8")
            result = run("build_report_skeleton.py", "--plan", str(root / "missing-plan.json"), "--evidence-pack", str(root / "missing-pack.json"), "--output", str(output))
            self.assertNotEqual(result.returncode, 0)
            self.assertEqual(output.read_text(encoding="utf-8"), "trusted previous report\n")

    def test_unknown_plan_field_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            value = json.loads(plan.read_text(encoding="utf-8"))
            value["unrequested_option"] = True
            plan.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            checked = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("unknown top-level", checked.stderr)

    def test_report_section_order_is_fixed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            value = json.loads(plan.read_text(encoding="utf-8"))
            value["quality_profile"] = "release"
            value["sections"][1], value["sections"][2] = value["sections"][2], value["sections"][1]
            plan.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            checked = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--final", "--root", str(root))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("sections must follow", checked.stderr)

    def test_question_title_in_pack_is_blocked_before_build(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            value = json.loads(evidence.read_text(encoding="utf-8"))
            value["analysis_points"][0]["title"] = "当前 GRN 优先哪些 TF-target 关系？"
            evidence.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            checked = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("declarative", checked.stderr)

    def test_unbounded_claim_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            value = json.loads(evidence.read_text(encoding="utf-8"))
            value["analysis_points"][0]["interpretation"] = "该结果证明了治疗效果"
            evidence.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            checked = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("overclaim", checked.stderr)

    def test_bounded_negative_claim_is_allowed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            value = json.loads(evidence.read_text(encoding="utf-8"))
            value["analysis_points"][0]["interpretation"] = "该关联不能证明因果"
            evidence.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            checked = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence))
            self.assertNotEqual(checked.returncode, 0, checked.stderr)
            self.assertIn("EVIDENCE_NEEDED", checked.stdout)

    def test_nonzero_docx_crop_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            docx = Path(temp) / "cropped.docx"
            xml = """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body><pic:pic><a:blip r:embed="rId1"/><a:srcRect l="50000"/></pic:pic></w:body></w:document>"""
            with zipfile.ZipFile(docx, "w") as archive:
                archive.writestr("word/document.xml", xml)
            checked = run("validate_docx_structure.py", str(docx))
            self.assertNotEqual(checked.returncode, 0)
            self.assertIn("srcRect", checked.stderr)

    def test_generic_builder_refuses_final_mode(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            docx = root / "report.docx"
            pack(evidence)
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            result = run("build_docx.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--output", str(docx), "--final")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("draft-only", result.stderr)
            self.assertFalse(docx.exists())

    def test_path_escape_is_blocked_when_root_is_supplied(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            evidence = root / "pack.json"
            plan = root / "plan.json"
            pack(evidence)
            value = json.loads(evidence.read_text(encoding="utf-8"))
            value["analysis_points"][0]["outputs"][0]["path"] = "../outside.csv"
            evidence.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            run("init_report_plan.py", "--evidence-pack", str(evidence), "--output", str(plan))
            result = run("validate_report_contract.py", "--plan", str(plan), "--evidence-pack", str(evidence), "--root", str(root))
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("relative", result.stderr)

    def test_docx_template_asset_contains_styled_note(self) -> None:
        result = run("validate_docx_structure.py", str(ROOT / "assets" / "report_template.docx"), "--json")
        self.assertNotEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["note_callouts"], 1)
        self.assertEqual(payload["note_style_issues"], [])
        self.assertEqual(payload["note_structure_issues"], [])
        self.assertEqual(payload["figure_box_issues"], [])
        self.assertEqual(payload["heading_style_issues"], [])
        self.assertIn("[[NOTE:DIRECTION]]", payload["template_markers"])

    def test_heading_run_cannot_override_title_style(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            docx = Path(temp) / "small-title.docx"
            from docx import Document
            from docx.shared import Pt

            document = Document()
            title = document.add_paragraph(style="Title")
            title_run = title.add_run("分析报告")
            title_run.font.size = Pt(11.5)
            title_run.bold = False
            document.save(docx)
            checked = run("validate_docx_structure.py", str(docx), "--final", "--json")
            self.assertNotEqual(checked.returncode, 0)
            payload = json.loads(checked.stdout)
            self.assertTrue(payload["heading_style_issues"])

    def test_docx_template_renderer_fills_note_tables_and_figure(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            (root / "result").mkdir()
            (root / "result" / "01.tsv").write_text("name\tvalue\nN\t4\n", encoding="utf-8")
            # A tiny valid PNG keeps this test independent of a plotting stack.
            png = base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            )
            (root / "result" / "01.png").write_bytes(png)
            values = {
                "quality_profile": "release",
                "slots": {
                    "REPORT_TITLE": "测试报告",
                    "REPORT_AUDIENCE": "研究人员",
                    "REPORT_SUMMARY": "摘要事实",
                    "ANALYSIS_SCOPE": "供体级样本",
                    "ANALYSIS_METHOD": "limma 3.58；FDR 0.05",
                    "ANALYSIS_QC": "样本顺序已核对",
                    "QC.CONDITION": "",
                    "ANALYSIS_RESULT": "观察到 4 个结果",
                    "RESULTS.CONDITION": "有结果",
                    "NOTE:DIRECTION": {"kind": "direction", "text": "logFC = case − control；正值表示 case 较高。"},
                    "RESULT_TABLE_ROWS": {"caption": "结果预览", "rows": [["N", 4, "gene", "result/01.tsv"]]},
                    "TABLE:RESULTS": "",
                    "ANALYSIS_CONCLUSION": "该结果为描述性证据",
                    "FIGURES.CONDITION": "有图",
                    "FIGURE:F1.TITLE": "差异表达火山图",
                    "FIGURE:F1.SOURCE": "result/01.png",
                    "FIGURE:F1.CAPTION": "图示供体级差异表达及比较方向。",
                    "FIGURE:F1.CAPTION_FIELDS": "对象；分组；轴和单位；统计层级；阈值；边界",
                    "ANALYSIS_LIMITATIONS": "样本量和外部验证范围有限",
                    "OUTPUTS_INTRO": "公开业务文件如下。",
                    "OUTPUT_TABLE_ROWS": {"caption": "公开文件", "rows": [
                        {"path": "result/01.tsv", "kind": "table", "description": "结果表", "purpose": "复核", "consumers": ["研究人员"]},
                        {"path": "report/report.docx", "kind": "report", "description": "内部报告", "purpose": "交付", "consumers": ["研究人员"]},
                        {"path": "log/run_record.json", "kind": "evidence", "description": "运行记录", "purpose": "审计", "consumers": ["审计"]},
                    ]},
                    "TABLE:OUTPUTS": "",
                    "REFERENCES": "limma 官方文档",
                    "VERSION_TABLE_ROWS": {"caption": "软件与资源版本", "rows": [["R", "4.3.0", "计算"]]},
                    "TABLE:VERSIONS": "",
                },
            }
            values_file = root / "values.json"
            values_file.write_text(json.dumps(values, ensure_ascii=False), encoding="utf-8")
            output = root / "report.docx"
            result = run(
                "render_docx_template.py",
                "--template", str(ROOT / "assets" / "report_template.docx"),
                "--values", str(values_file), "--root", str(root), "--output", str(output),
                "--final", "--require-note", "--json",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual(payload["status"], "PASS")
            self.assertEqual(payload["unresolved"], [])
            checked = run("validate_docx_structure.py", str(output), "--final", "--require-note", "--json")
            self.assertEqual(checked.returncode, 0, checked.stdout + checked.stderr)
            rendered = json.loads(checked.stdout)
            self.assertEqual(rendered["drawings"], 1)
            self.assertEqual(rendered["note_style_issues"], [])
            rendered_document = __import__("docx").Document(output)
            document_text = "\n".join(
                [paragraph.text for paragraph in rendered_document.paragraphs]
                + [cell.text for table in rendered_document.tables for row in table.rows for cell in row.cells]
            )
            self.assertNotIn("[[", document_text)
            self.assertNotIn("条件图件", document_text)
            self.assertNotIn("图注字段（由 renderer", document_text)
            self.assertNotIn("TRUE", document_text)
            self.assertIn("result/01.tsv", document_text)
            self.assertNotIn("report/report.docx", document_text)
            self.assertNotIn("run_record.json", document_text)

    def test_docx_template_renderer_keeps_source_urls_out_of_visible_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            values = {
                "quality_profile": "release",
                "slots": {
                    "REPORT_TITLE": "来源报告", "REPORT_AUDIENCE": "研究人员", "REPORT_SUMMARY": "摘要",
                    "ANALYSIS_SCOPE": "范围", "ANALYSIS_METHOD": "方法", "ANALYSIS_RESULT": "结果",
                    "ANALYSIS_CONCLUSION": "结论", "ANALYSIS_LIMITATIONS": "限制", "OUTPUTS_INTRO": "公开文件",
                    "REFERENCES": [{"name": "limma", "version": "3.58", "source": "https://example.org/limma", "purpose": "统计方法"}],
                    "RESULT_TABLE_ROWS": [], "OUTPUT_TABLE_ROWS": [], "VERSION_TABLE_ROWS": [],
                },
            }
            values_file = root / "values.json"
            values_file.write_text(json.dumps(values, ensure_ascii=False), encoding="utf-8")
            output = root / "report.docx"
            result = run(
                "render_docx_template.py", "--template", str(ROOT / "assets" / "report_template.docx"),
                "--values", str(values_file), "--output", str(output), "--final", "--json",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            document = __import__("docx").Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertNotIn("https://example.org", text)
            self.assertIn("limma", text)

    def test_docx_template_renderer_removes_optional_qc_note_and_figures(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            values = {
                "slots": {
                    "REPORT_TITLE": "简版报告",
                    "REPORT_AUDIENCE": "研究人员",
                    "REPORT_SUMMARY": "结果摘要",
                    "ANALYSIS_SCOPE": "样本范围",
                    "ANALYSIS_METHOD": "方法与参数",
                    "ANALYSIS_RESULT": "结果段",
                    "ANALYSIS_CONCLUSION": "结论",
                    "ANALYSIS_LIMITATIONS": "限制",
                    "OUTPUTS_INTRO": "公开文件",
                    "REFERENCES": "实际来源",
                    "RESULT_TABLE_ROWS": [],
                    "OUTPUT_TABLE_ROWS": [],
                    "VERSION_TABLE_ROWS": [],
                }
            }
            values_file = root / "values.json"
            values_file.write_text(json.dumps(values, ensure_ascii=False), encoding="utf-8")
            output = root / "report.docx"
            result = run(
                "render_docx_template.py",
                "--template", str(ROOT / "assets" / "report_template.docx"),
                "--values", str(values_file), "--root", str(root), "--output", str(output), "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn('"status": "DRAFT"', result.stdout)
            document = __import__("docx").Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertNotIn("质控与异常", text)
            self.assertNotIn("图件", text)
            self.assertEqual(len(document.tables), 3)  # result/output/version; Note is conditional

    def test_docx_template_renderer_keeps_multiple_explicit_notes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            values = {
                "slots": {
                    "REPORT_TITLE": "双提示报告", "REPORT_AUDIENCE": "研究人员", "REPORT_SUMMARY": "摘要",
                    "ANALYSIS_SCOPE": "范围", "ANALYSIS_METHOD": "方法", "ANALYSIS_RESULT": "结果",
                    "ANALYSIS_CONCLUSION": "结论", "ANALYSIS_LIMITATIONS": "限制", "OUTPUTS_INTRO": "公开文件",
                    "REFERENCES": "来源", "RESULT_TABLE_ROWS": [], "OUTPUT_TABLE_ROWS": [], "VERSION_TABLE_ROWS": [],
                    "notes": [
                        {"id": "direction", "kind": "direction", "text": "logFC = case − control。"},
                        {"id": "unit", "kind": "unit", "text": "统计单位为供体。"},
                    ],
                }
            }
            values_file = root / "values.json"
            values_file.write_text(json.dumps(values, ensure_ascii=False), encoding="utf-8")
            output = root / "report.docx"
            result = run(
                "render_docx_template.py", "--template", str(ROOT / "assets" / "report_template.docx"),
                "--values", str(values_file), "--output", str(output), "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            document = __import__("docx").Document(output)
            note_tables = [
                table for table in document.tables
                if "Note：" in "".join(cell.text for row in table.rows for cell in row.cells)
            ]
            self.assertEqual(len(note_tables), 2)
            note_text = "\n".join(cell.text for table in note_tables for row in table.rows for cell in row.cells)
            self.assertIn("统计单位为供体", note_text)

    def test_legacy_figure_box_is_cleared_and_wide_image_fits_body(self) -> None:
        """旧模板的图段蓝框须清除，宽图按页面版心等比缩放。"""
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            legacy = root / "legacy.docx"
            shutil.copy2(ROOT / "assets" / "report_template.docx", legacy)
            from PIL import Image
            image = root / "wide.png"
            Image.new("RGB", (2400, 400), "white").save(image)
            from docx import Document
            from docx.oxml import OxmlElement
            document = Document(legacy)
            source = next(
                paragraph for paragraph in document.paragraphs
                if "[[FIGURE:F1.SOURCE]]" in paragraph.text
            )
            ppr = source._p.get_or_add_pPr()
            for name in ("w:pBdr", "w:shd", "w:ind"):
                ppr.append(OxmlElement(name))
            # 旧模板可能已经含有一个 drawing；renderer 必须替换它而不是叠加。
            source.add_run().add_picture(str(image))
            document.save(legacy)
            values = {
                "quality_profile": "release",
                "slots": {
                    "REPORT_TITLE": "宽图报告", "REPORT_AUDIENCE": "研究人员", "REPORT_SUMMARY": "摘要",
                    "ANALYSIS_SCOPE": "范围", "ANALYSIS_METHOD": "方法", "ANALYSIS_QC": "质控",
                    "ANALYSIS_RESULT": "结果", "ANALYSIS_CONCLUSION": "结论", "ANALYSIS_LIMITATIONS": "限制",
                    "NOTE:DIRECTION": {"kind": "direction", "text": "case − control"},
                    "RESULT_TABLE_ROWS": [], "OUTPUT_TABLE_ROWS": [], "VERSION_TABLE_ROWS": [],
                    "FIGURE:F1.TITLE": "宽图", "FIGURE:F1.SOURCE": "wide.png",
                    "FIGURE:F1.CAPTION": "对象、分组、轴、单位、统计层级、边界",
                    "OUTPUTS_INTRO": "公开文件", "REFERENCES": "来源",
                },
            }
            values_file = root / "values.json"
            values_file.write_text(json.dumps(values, ensure_ascii=False), encoding="utf-8")
            output = root / "report.docx"
            result = run(
                "render_docx_template.py", "--template", str(legacy), "--values", str(values_file),
                "--root", str(root), "--output", str(output), "--final", "--require-note", "--json",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual(payload["status"], "PASS")
            self.assertEqual(payload["unresolved"], [])
            self.assertEqual(payload.get("figure_box_issues", []), [])
            rendered = Document(output)
            self.assertEqual(len(rendered.inline_shapes), 1)
            body_width = rendered.sections[0].page_width - rendered.sections[0].left_margin - rendered.sections[0].right_margin
            shape = rendered.inline_shapes[0]
            self.assertLessEqual(shape.width, body_width)
            self.assertAlmostEqual(shape.width / shape.height, 6.0, places=3)
            checked = run("validate_docx_structure.py", str(output), "--final", "--require-note", "--json")
            self.assertEqual(checked.returncode, 0, checked.stdout + checked.stderr)
            self.assertEqual(json.loads(checked.stdout).get("figure_box_issues", []), [])

    def test_failed_final_template_render_does_not_replace_existing_report(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            target = root / "report.docx"
            target.write_bytes(b"previous release")
            values = root / "values.json"
            values.write_text(json.dumps({"quality_profile": "release", "slots": {}}, ensure_ascii=False), encoding="utf-8")
            result = run(
                "render_docx_template.py",
                "--template", str(ROOT / "assets" / "report_template.docx"),
                "--values", str(values), "--root", str(root), "--output", str(target), "--final", "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertEqual(target.read_bytes(), b"previous release")


if __name__ == "__main__":
    unittest.main()
